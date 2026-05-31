"""
tests/test_exam_router.py
════════════════════════════════════════════════════════════════════════════
Layer 2: Integration tests for POST /exams/parse
Tests the full HTTP layer using FastAPI's TestClient.
Auth is bypassed by overriding the get_current_teacher dependency.
════════════════════════════════════════════════════════════════════════════
"""

import io
import os
import pytest
from fastapi.testclient import TestClient
from docx import Document

# ─── App & Dependency Override Setup ─────────────────────────────────────────
# Import the FastAPI app. Adjust path if your app factory is elsewhere.
import sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from main import app
from core.dependencies import get_current_teacher
from core.config import settings


# Mock teacher object — same shape the dependency normally returns
class _MockTeacher:
    id = 1
    email = "teacher@test.com"
    is_teacher = True


def _override_teacher():
    return _MockTeacher()


app.dependency_overrides[get_current_teacher] = _override_teacher

client = TestClient(app)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _make_docx_bytes(questions: list[dict]) -> bytes:
    """Build a DOCX in-memory and return its raw bytes.

    Each item in `questions` is a dict:
        { "stem": str, "options": [str, ...], "bold_index": int | None }
    """
    doc = Document()
    labels = ["A", "B", "C", "D"]
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q['stem']}")
        for j, opt in enumerate(q["options"]):
            p = doc.add_paragraph()
            run = p.add_run(f"{labels[j]}) {opt}")
            if q.get("bold_index") == j:
                run.bold = True
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _post_docx(content: bytes, filename: str = "exam.docx") -> dict:
    """POST content to /exams/parse and return the response."""
    return client.post(
        f"{settings.API_PREFIX}/exams/parse",
        files={"file": (filename, content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )


# ─── R01: Happy Path ──────────────────────────────────────────────────────────

def test_R01_happy_path_returns_200_with_questions():
    """Valid DOCX with 3 fully bolded answers → 200, 3 questions, no warnings."""
    qs = [
        {"stem": "What is 2+2?", "options": ["3", "4", "5", "6"], "bold_index": 1},
        {"stem": "Sky color?", "options": ["Red", "Blue", "Green", "Yellow"], "bold_index": 1},
        {"stem": "Capital of Egypt?", "options": ["Cairo", "London", "Paris", "Berlin"], "bold_index": 0},
    ]
    resp = _post_docx(_make_docx_bytes(qs))

    assert resp.status_code == 200
    body = resp.json()
    assert len(body["questions"]) == 3
    assert body["warnings"] == []
    for q in body["questions"]:
        assert q["marks"] == 1
        assert q["correct_answer"] is not None


# ─── R02: Partial Success (missing answer) ────────────────────────────────────

def test_R02_missing_answer_returns_200_with_warnings():
    """DOCX with Q2 missing bold → 200, Q2.correct_answer is null, warnings populated."""
    qs = [
        {"stem": "Q1?", "options": ["A", "B", "C", "D"], "bold_index": 0},
        {"stem": "Q2 no answer?", "options": ["A", "B", "C", "D"], "bold_index": None},
        {"stem": "Q3?", "options": ["A", "B", "C", "D"], "bold_index": 2},
    ]
    resp = _post_docx(_make_docx_bytes(qs))

    assert resp.status_code == 200
    body = resp.json()
    assert body["questions"][1]["correct_answer"] is None
    assert len(body["warnings"]) == 1


# ─── R03: Non-MCQ Document ───────────────────────────────────────────────────

def test_R03_non_mcq_document_returns_400():
    """Plain text syllabus with no numbered questions → 400 with parser message."""
    doc = Document()
    doc.add_paragraph("Week 1: Introduction")
    doc.add_paragraph("Week 2: Advanced topics")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    resp = _post_docx(buf.read())

    assert resp.status_code == 400
    assert "No valid multiple-choice questions" in resp.json()["detail"]


# ─── R04: Empty File ──────────────────────────────────────────────────────────

def test_R04_empty_file_returns_400():
    """0-byte upload → 400 with 'Uploaded file is empty' message."""
    resp = _post_docx(b"", filename="empty.docx")

    assert resp.status_code == 400
    assert "empty" in resp.json()["detail"].lower()


# ─── R05: Oversized File ──────────────────────────────────────────────────────

def test_R05_oversized_file_returns_413():
    """Upload > 10 MB → 413 with size limit message."""
    big_content = b"x" * (10 * 1024 * 1024 + 1)  # 10 MB + 1 byte
    resp = _post_docx(big_content, filename="big.docx")

    assert resp.status_code == 413
    assert "10 MB" in resp.json()["detail"]


# ─── R06: Wrong File Type ─────────────────────────────────────────────────────

def test_R06_wrong_file_type_returns_415():
    """PNG file uploaded → 415 with unsupported type message."""
    resp = client.post(
        f"{settings.API_PREFIX}/exams/parse",
        files={"file": ("scan.png", b"\x89PNG\r\n", "image/png")},
    )

    assert resp.status_code == 415
    assert ".png" in resp.json()["detail"]


# ─── R07: No Filename ────────────────────────────────────────────────────────

def test_R07_no_filename_returns_422():
    """Upload with empty filename → 422 from FastAPI validation."""
    resp = client.post(
        f"{settings.API_PREFIX}/exams/parse",
        files={"file": ("", b"some bytes", "application/octet-stream")},
    )

    assert resp.status_code == 422


# ─── R08: Unauthenticated Request ────────────────────────────────────────────

def test_R08_unauthenticated_request_returns_401():
    """Request with no auth token → 401. Use a fresh client with no overrides."""
    # Create a fresh app instance without the dependency override
    from fastapi.testclient import TestClient as FreshClient
    fresh_app_client = FreshClient(app)

    # Temporarily remove the override
    app.dependency_overrides.pop(get_current_teacher, None)
    try:
        resp = fresh_app_client.post(
            f"{settings.API_PREFIX}/exams/parse",
            files={"file": ("exam.docx", b"content", "application/octet-stream")},
        )
        assert resp.status_code in (401, 403)
    finally:
        # Restore override so other tests are not affected
        app.dependency_overrides[get_current_teacher] = _override_teacher


# ─── R09: Questions Schema Validation ────────────────────────────────────────

def test_R09_response_schema_is_correct():
    """Verify the exact JSON shape returned matches the ParsedExamResult contract."""
    qs = [{"stem": "Test?", "options": ["A", "B", "C", "D"], "bold_index": 1}]
    resp = _post_docx(_make_docx_bytes(qs))

    assert resp.status_code == 200
    body = resp.json()

    # Top-level keys
    assert "questions" in body
    assert "warnings" in body

    # Question shape
    q = body["questions"][0]
    assert "number" in q
    assert "text" in q
    assert "marks" in q
    assert "options" in q
    assert "correct_answer" in q

    # No stale fields from old models
    assert "detection_method" not in q
    assert "question_count" not in body
    assert "source_file" not in body
    assert "has_fractional_error" not in body


# ─── R10: Temp File Cleanup ──────────────────────────────────────────────────

def test_R10_temp_file_is_deleted_after_parse(monkeypatch):
    """After a successful parse, no temp file should remain on disk."""
    created_paths = []
    original_named_temp = __import__("tempfile").NamedTemporaryFile

    import tempfile

    def tracking_named_temp(*args, **kwargs):
        tmp = original_named_temp(*args, **kwargs)
        created_paths.append(tmp.name)
        return tmp

    monkeypatch.setattr(tempfile, "NamedTemporaryFile", tracking_named_temp)

    qs = [{"stem": "Clean up?", "options": ["Yes", "No", "Maybe", "Never"], "bold_index": 0}]
    _post_docx(_make_docx_bytes(qs))

    for path in created_paths:
        assert not os.path.exists(path), f"Temp file was NOT cleaned up: {path}"
