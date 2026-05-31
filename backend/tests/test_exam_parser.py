"""
tests/test_exam_parser.py
════════════════════════════════════════════════════════════════════════════
Layer 1: Unit tests for services/exam_parser.py
Tests the parse_exam() service directly, with no HTTP or FastAPI involved.
All DOCX fixtures are built programmatically using python-docx.
════════════════════════════════════════════════════════════════════════════
"""

import os
import pytest
from pathlib import Path
from docx import Document
from services.exam_parser import parse_exam, ParsedExamResult


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _save_docx(doc: Document, tmp_path: Path, name: str) -> Path:
    """Save a Document to the pytest temp directory and return the path."""
    path = tmp_path / name
    doc.save(str(path))
    return path


def _make_question(doc: Document, num: int, stem: str,
                   options: list[str], bold_index: int | None):
    """Append a numbered MCQ to a Document.

    bold_index: index of the option to mark bold (correct answer).
                Pass None to produce a question with no bold answer.
    """
    doc.add_paragraph(f"{num}. {stem}")
    labels = ["A", "B", "C", "D"]
    for i, opt in enumerate(options):
        p = doc.add_paragraph()
        run = p.add_run(f"{labels[i]}) {opt}")
        if bold_index is not None and i == bold_index:
            run.bold = True


# ─── T01: Happy Path ──────────────────────────────────────────────────────────

def test_T01_happy_path_all_answers_bolded(tmp_path):
    """All 3 questions have a bolded correct answer. Should return clean result."""
    doc = Document()
    _make_question(doc, 1, "What is 2+2?", ["3", "4", "5", "6"], bold_index=1)
    _make_question(doc, 2, "What color is the sky?", ["Red", "Green", "Blue", "Yellow"], bold_index=2)
    _make_question(doc, 3, "Capital of Egypt?", ["Cairo", "London", "Paris", "Berlin"], bold_index=0)
    path = _save_docx(doc, tmp_path, "happy.docx")

    result = parse_exam(path)

    assert isinstance(result, ParsedExamResult)
    assert len(result.questions) == 3
    assert result.warnings == []
    assert result.questions[0].correct_answer == "4"
    assert result.questions[1].correct_answer == "Blue"
    assert result.questions[2].correct_answer == "Cairo"


# ─── T02: Partial Success (one missing answer) ────────────────────────────────

def test_T02_one_missing_answer(tmp_path):
    """Q2 has no bold answer. Should return 3 questions, Q2.correct_answer is None."""
    doc = Document()
    _make_question(doc, 1, "Q1 stem?", ["A", "B", "C", "D"], bold_index=0)
    _make_question(doc, 2, "Q2 stem?", ["A", "B", "C", "D"], bold_index=None)
    _make_question(doc, 3, "Q3 stem?", ["A", "B", "C", "D"], bold_index=3)
    path = _save_docx(doc, tmp_path, "partial.docx")

    result = parse_exam(path)

    assert len(result.questions) == 3
    assert result.questions[0].correct_answer is not None
    assert result.questions[1].correct_answer is None
    assert result.questions[2].correct_answer is not None
    assert len(result.warnings) == 1
    assert "2" in result.warnings[0]


# ─── T03: All Answers Missing ─────────────────────────────────────────────────

def test_T03_all_answers_missing(tmp_path):
    """No questions have bold answers. Warnings should list all question numbers."""
    doc = Document()
    for i in range(1, 4):
        _make_question(doc, i, f"Question {i} stem?", ["A", "B", "C", "D"], bold_index=None)
    path = _save_docx(doc, tmp_path, "no_answers.docx")

    result = parse_exam(path)

    assert len(result.questions) == 3
    for q in result.questions:
        assert q.correct_answer is None
    assert len(result.warnings) == 1
    assert "1" in result.warnings[0]
    assert "2" in result.warnings[0]
    assert "3" in result.warnings[0]


# ─── T04: Non-MCQ Document (zero questions) ───────────────────────────────────

def test_T04_non_mcq_document_raises(tmp_path):
    """A plain paragraph doc with no numbered questions raises ValueError."""
    doc = Document()
    doc.add_paragraph("This is a syllabus.")
    doc.add_paragraph("Week 1: Introduction to the course.")
    doc.add_paragraph("Week 2: Advanced topics.")
    path = _save_docx(doc, tmp_path, "syllabus.docx")

    with pytest.raises(ValueError, match="No valid multiple-choice questions found"):
        parse_exam(path)


# ─── T05: Empty Document ──────────────────────────────────────────────────────

def test_T05_empty_document_raises(tmp_path):
    """An empty DOCX (no paragraphs) raises ValueError."""
    doc = Document()
    path = _save_docx(doc, tmp_path, "empty.docx")

    with pytest.raises(ValueError, match="No valid multiple-choice questions found"):
        parse_exam(path)


# ─── T06: Marks Are Always 1 ──────────────────────────────────────────────────

def test_T06_marks_always_1(tmp_path):
    """Every extracted question must have marks == 1, regardless of content."""
    doc = Document()
    for i in range(1, 6):
        _make_question(doc, i, f"Question {i}?", ["Opt A", "Opt B", "Opt C", "Opt D"], bold_index=0)
    path = _save_docx(doc, tmp_path, "five_questions.docx")

    result = parse_exam(path)

    assert len(result.questions) == 5
    for q in result.questions:
        assert q.marks == 1


# ─── T07: Multi-line Question Stem ────────────────────────────────────────────

def test_T07_multiline_stem(tmp_path):
    """A question stem that spans 2 paragraphs should be concatenated."""
    doc = Document()
    doc.add_paragraph("1. This is the first line of the question,")
    doc.add_paragraph("and this is the continuation of the stem.")
    doc.add_paragraph("A) Option 1")
    p = doc.add_paragraph()
    run = p.add_run("B) Option 2")
    run.bold = True
    doc.add_paragraph("C) Option 3")
    doc.add_paragraph("D) Option 4")
    path = _save_docx(doc, tmp_path, "multiline.docx")

    result = parse_exam(path)

    assert len(result.questions) == 1
    assert "continuation" in result.questions[0].text


# ─── T08: Different Option Delimiters ────────────────────────────────────────

def test_T08_option_delimiter_variants(tmp_path):
    """Options formatted as A., A), (A), or [A] must all be parsed correctly."""
    doc = Document()
    doc.add_paragraph("1. Which delimiter works?")
    doc.add_paragraph("A. First")
    p = doc.add_paragraph()
    p.add_run("B) Second").bold = True
    doc.add_paragraph("(C) Third")
    doc.add_paragraph("[D] Fourth")
    path = _save_docx(doc, tmp_path, "delimiters.docx")

    result = parse_exam(path)

    assert len(result.questions) == 1
    assert len(result.questions[0].options) == 4
    assert result.questions[0].correct_answer == "Second"


# ─── T09: Decorative Bold Label Only ─────────────────────────────────────────

def test_T09_decorative_bold_label_is_not_correct_answer(tmp_path):
    """If only the option label letter is bold (e.g. **A)**), it should NOT be
    treated as the correct answer — only the content being bold counts."""
    doc = Document()
    doc.add_paragraph("1. Which is correct?")

    # Label is bold but option text is not
    p = doc.add_paragraph()
    run_label = p.add_run("A) ")
    run_label.bold = True
    p.add_run("This is Option A")

    doc.add_paragraph("B) This is Option B")
    doc.add_paragraph("C) This is Option C")
    p2 = doc.add_paragraph()
    p2.add_run("D) This is Option D").bold = True  # entire option bold = correct
    path = _save_docx(doc, tmp_path, "decorative_bold.docx")

    result = parse_exam(path)

    assert len(result.questions) == 1
    # A's label-only bold should not mark it correct
    # D's full bold should mark it correct
    assert result.questions[0].correct_answer == "This is Option D"


# ─── T10: Unsupported Extension ───────────────────────────────────────────────

def test_T10_unsupported_extension_raises(tmp_path):
    """Passing a .txt file path raises ValueError about unsupported type."""
    path = tmp_path / "exam.txt"
    path.write_text("Some text content")

    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_exam(path)


# ─── T11: Image Extension ─────────────────────────────────────────────────────

def test_T11_image_extension_raises(tmp_path):
    """Passing a .png path raises ValueError about image files."""
    path = tmp_path / "scan.png"
    path.write_bytes(b"\x89PNG\r\n")  # fake PNG header

    with pytest.raises(ValueError, match="Image files cannot be used"):
        parse_exam(path)


# ─── T12: Non-Existent File ───────────────────────────────────────────────────

def test_T12_nonexistent_file_raises(tmp_path):
    """Passing a path that doesn't exist raises FileNotFoundError."""
    path = tmp_path / "ghost.docx"

    with pytest.raises(FileNotFoundError):
        parse_exam(path)
