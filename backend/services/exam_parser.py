"""
exam_parser.py
══════════════════════════════════════════════════════════════════════════════
Local MCQ Exam Parser — single-file service, no LLM, no internet required.

SUPPORTS : .pdf, .docx
DETECTS  : questions · MCQ options · correct answer (bold option)

USAGE
─────
    from services.exam_parser import parse_exam

    result = parse_exam("midterm_calculus.pdf")
    print(result.summary())
══════════════════════════════════════════════════════════════════════════════
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from pydantic import BaseModel, Field


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 · MODELS
# ══════════════════════════════════════════════════════════════════════════════

class Question(BaseModel):
    """A single MCQ question with its options and detected answer."""
    number: int = Field(..., description="Question number as it appears in the exam")
    text: str = Field(..., description="The question stem text")
    marks: int = Field(..., description="Marks allocated to this question")
    options: list[str] = Field(default_factory=list, description="List of option strings")
    correct_answer: Optional[str] = Field(None, description="Exact text of the correct option")

    @property
    def has_answer(self) -> bool:
        return self.correct_answer is not None


class ParsedExamResult(BaseModel):
    """Full parsed exam result — top-level output of parse_exam()."""
    questions: list[Question] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list, description="Non-fatal issues during parsing")

    def summary(self) -> str:
        """Human-readable summary of the parsed exam."""
        answered = sum(1 for q in self.questions if q.has_answer)
        lines = [f"Questions: {len(self.questions)}, Answered: {answered}/{len(self.questions)}"]
        if self.warnings:
            lines.append("Warnings:")
            for w in self.warnings:
                lines.append(f"  {w}")
        return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 · CORE TYPES
# ══════════════════════════════════════════════════════════════════════════════

class TextRun:
    """Smallest unit of text with a bold flag."""
    __slots__ = ("text", "bold")

    def __init__(self, text: str, bold: bool = False):
        self.text = text
        self.bold = bold

    def __repr__(self):
        return f"TextRun({self.text!r}, bold={self.bold})"


class Line:
    """A logical line composed of one or more TextRuns."""
    def __init__(self, runs: list[TextRun]):
        self.runs = runs

    @property
    def full_text(self) -> str:
        return "".join(r.text for r in self.runs).strip()

    @property
    def is_fully_bold(self) -> bool:
        text_runs = [r for r in self.runs if r.text.strip()]
        return bool(text_runs) and all(r.bold for r in text_runs)

    @property
    def has_any_bold(self) -> bool:
        return any(r.bold and r.text.strip() for r in self.runs)

    def bold_text(self) -> str:
        return "".join(r.text for r in self.runs if r.bold).strip()

    def __repr__(self):
        return f"Line({self.full_text!r}, fully_bold={self.is_fully_bold})"


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 · COMPILED REGEX PATTERNS
# ══════════════════════════════════════════════════════════════════════════════

# Question start: "Q1.", "Question 1", "1.", "1)"
_QUESTION_START_RE = re.compile(
    r"^(?:Q(?:uestion)?\s*)?(\d+)[.)]\s*(.+)",
    re.IGNORECASE,
)

# MCQ option line: "A. text", "A) text", "(A) text", "a. text"
_OPTION_RE = re.compile(
    r"^\s*[\[(]?\s*([A-Da-d])\s*[.):\]]\s*(.+)",
)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 · QUESTION / OPTION / ANSWER PARSER
# ══════════════════════════════════════════════════════════════════════════════

def _parse_option_line(line: Line) -> Optional[tuple[str, str, bool]]:
    """Try to interpret a Line as an MCQ option. Returns (label_upper, option_text, is_correct) or None."""
    m = _OPTION_RE.match(line.full_text)
    if not m:
        return None

    label = m.group(1).upper()
    option_text = m.group(2).strip()

    if line.is_fully_bold:
        return label, option_text, True

    if line.has_any_bold:
        bold_part = line.bold_text().strip().rstrip(".):")
        # If the bold portion is only the single label letter → decorative, skip
        if bold_part.lower() == label.lower():
            return label, option_text, False
        return label, option_text, True

    return label, option_text, False


def _parse_questions(lines: list[Line]) -> list[Question]:
    questions: list[Question] = []
    current_q: Optional[dict] = None
    state = "SCANNING"

    def flush():
        nonlocal current_q
        if current_q is None:
            return
        options_text = [o["text"] for o in current_q["options"]]
        correct = next((o["text"] for o in current_q["options"] if o["is_correct"]), None)
        
        questions.append(Question(
            number=current_q["number"],
            text=current_q["text"],
            marks=1,
            options=options_text,
            correct_answer=correct,
        ))
        current_q = None

    for line in lines:
        text = line.full_text
        if not text:
            continue

        # New question?
        q_match = _QUESTION_START_RE.match(text)
        if q_match:
            flush()
            q_num = int(q_match.group(1))
            raw_stem = q_match.group(2).strip()
            current_q = {
                "number": q_num,
                "text": raw_stem,
                "options": [],
            }
            state = "IN_QUESTION"
            continue

        # MCQ option?
        if state in ("IN_QUESTION", "IN_OPTIONS"):
            parsed = _parse_option_line(line)
            if parsed and current_q is not None:
                label, opt_text, is_correct = parsed
                current_q["options"].append({
                    "label": label,
                    "text": opt_text,
                    "is_correct": is_correct,
                })
                state = "IN_OPTIONS"
                continue

        # Continuation of multi-line question stem
        if state == "IN_QUESTION" and current_q is not None:
            current_q["text"] += " " + text

    flush()

    return questions


def _parse_structure(lines: list[Line]) -> ParsedExamResult:
    warnings: list[str] = []
    
    questions = _parse_questions(lines)
    
    if not questions:
        raise ValueError(
            "No valid multiple-choice questions found. "
            "Please ensure the document contains numbered questions and lettered choices (A, B, C, D)."
        )

    no_answer = [q.number for q in questions if not q.has_answer]
    if no_answer:
        warnings.append(f"No bold answer detected for question(s): {no_answer}.")

    return ParsedExamResult(
        questions=questions,
        warnings=warnings,
    )


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 · PDF EXTRACTOR
# ══════════════════════════════════════════════════════════════════════════════

# PyMuPDF encodes bold as bit 4 in the font flags integer
_PDF_BOLD_FLAG = 1 << 4  # = 16

def _is_bold_by_fontname(font_name: str) -> bool:
    name_lower = font_name.lower()
    return any(kw in name_lower for kw in ("bold", "heavy", "black", "demi"))

def _extract_lines_from_pdf(file_path: Path) -> list[Line]:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError(
            "PyMuPDF is required for PDF parsing. "
            "Install it with: pip install pymupdf"
        )

    doc = fitz.open(str(file_path))
    all_lines: list[Line] = []

    for page in doc:
        raw = page.get_text("dict")
        for block in raw.get("blocks", []):
            if block.get("type") != 0:
                continue
            for raw_line in block.get("lines", []):
                runs: list[TextRun] = []
                for span in raw_line.get("spans", []):
                    span_text: str = span.get("text", "")
                    if not span_text:
                        continue
                    flags: int = span.get("flags", 0)
                    font: str = span.get("font", "")
                    is_bold = bool(flags & _PDF_BOLD_FLAG) or _is_bold_by_fontname(font)
                    runs.append(TextRun(text=span_text, bold=is_bold))
                if runs:
                    line = Line(runs)
                    if line.full_text.strip():
                        all_lines.append(line)

    doc.close()
    return all_lines


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 · DOCX EXTRACTOR
# ══════════════════════════════════════════════════════════════════════════════

def _resolve_bold_docx(run) -> bool:
    if run.bold is True:
        return True
    if run.bold is False:
        return False
    style = run.style
    while style is not None:
        try:
            if style.font.bold is True:
                return True
            if style.font.bold is False:
                return False
        except AttributeError:
            # Malformed or non-standard style with no font object — treat as non-bold
            break
        style = style.base_style
    return False

def _extract_lines_from_docx(file_path: Path) -> list[Line]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install it with: pip install python-docx"
        )

    doc = Document(str(file_path))
    all_lines: list[Line] = []

    def _para_to_line(para) -> Optional[Line]:
        if not para.text.strip():
            return None
        if not para.runs:
            return Line([TextRun(text=para.text, bold=False)])
        runs = [
            TextRun(text=r.text, bold=_resolve_bold_docx(r))
            for r in para.runs if r.text
        ]
        return Line(runs) if runs else None

    # Main paragraphs
    for para in doc.paragraphs:
        line = _para_to_line(para)
        if line and line.full_text.strip():
            all_lines.append(line)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    line = _para_to_line(para)
                    if line and line.full_text.strip():
                        all_lines.append(line)

    return all_lines


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 · FILE ROUTER & PUBLIC API
# ══════════════════════════════════════════════════════════════════════════════

_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"}

def _route(file_path: Path) -> list[Line]:
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = file_path.suffix.lower()

    if suffix in _IMAGE_EXTENSIONS:
        raise ValueError("Image files cannot be used for bold-based answer detection.")
    if suffix not in _SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: '{suffix}'.")

    if suffix == ".pdf":
        return _extract_lines_from_pdf(file_path)
    if suffix in (".docx", ".doc"):
        return _extract_lines_from_docx(file_path)

    raise ValueError(f"No extractor registered for extension: {suffix}")

def parse_exam(file_path: str | Path) -> ParsedExamResult:
    """Parse an MCQ exam file and return a structured ParsedExamResult."""
    path = Path(file_path)
    lines = _route(path)
    return _parse_structure(lines)
